from __future__ import annotations

import math
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "Urban_Thermal_Math_Deep_Dive_Picturized.docx"
TMP_DIR = Path(tempfile.mkdtemp(prefix="picturized_deep_dive_", dir="/private/tmp"))

ASSETS = {
    "boston_map": ROOT / "boston_heatmap.png",
    "fig1": ROOT / "fig1.png",
    "fig2": ROOT / "fig2.png",
    "fig3": ROOT / "fig3.png",
    "fig4": ROOT / "fig4.png",
}

NAVY = "#123B6D"
BLUE = "#2E74B5"
SKY = "#DCEBFA"
INK = "#1F2937"
MUTED = "#5B6472"
BG = "#F7F8FA"
CARD = "#FFFFFF"
GREEN = "#2E8B57"
GOLD = "#C47F00"
RED = "#B3261E"
PURPLE = "#7A4CC2"
TEAL = "#0F766E"
LIGHT_GREEN = "#E7F3EC"
LIGHT_GOLD = "#FFF3D7"
LIGHT_RED = "#FDECEC"
LIGHT_PURPLE = "#F1E9FF"

FONT_DIR = Path("/System/Library/Fonts/Supplemental")
FONT_REG = FONT_DIR / "Arial.ttf"
FONT_BOLD = FONT_DIR / "Arial Bold.ttf"
FONT_ITALIC = FONT_DIR / "Arial Italic.ttf"


def font(size: int, bold: bool = False, italic: bool = False):
    path = FONT_REG
    if bold and italic:
        path = FONT_DIR / "Arial Bold Italic.ttf"
    elif bold:
        path = FONT_BOLD
    elif italic:
        path = FONT_ITALIC
    return ImageFont.truetype(str(path), size=size)


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont):
    bbox = draw.textbbox((0, 0), text, font=fnt)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.FreeTypeFont, max_width: int):
    words = text.split()
    lines = []
    current = []
    for word in words:
        trial = " ".join(current + [word])
        w, _ = text_size(draw, trial, fnt)
        if w <= max_width or not current:
            current.append(word)
        else:
            lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines


def draw_wrapped(draw, xy, text, fnt, fill, max_width, line_gap=10):
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    for line in lines:
        draw.text((x, y), line, font=fnt, fill=fill)
        _, h = text_size(draw, line, fnt)
        y += h + line_gap
    return y


def rounded_box(draw, box, fill, outline=None, width=2, radius=28):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color, width=8):
    draw.line([start, end], fill=color, width=width)
    ang = math.atan2(end[1] - start[1], end[0] - start[0])
    ah = 18
    aw = 10
    p1 = end
    p2 = (end[0] - ah * math.cos(ang) + aw * math.sin(ang), end[1] - ah * math.sin(ang) - aw * math.cos(ang))
    p3 = (end[0] - ah * math.cos(ang) - aw * math.sin(ang), end[1] - ah * math.sin(ang) + aw * math.cos(ang))
    draw.polygon([p1, p2, p3], fill=color)


def circle_label(draw, center, r, fill, outline=None, text="", text_fill="#FFFFFF", fnt=None):
    x, y = center
    draw.ellipse((x - r, y - r, x + r, y + r), fill=fill, outline=outline, width=2)
    if text:
        if fnt is None:
            fnt = font(24, bold=True)
        tw, th = text_size(draw, text, fnt)
        draw.text((x - tw / 2, y - th / 2 - 2), text, font=fnt, fill=text_fill)


def create_canvas(title: str, subtitle: str = "", theme=BLUE):
    img = Image.new("RGB", (1600, 900), BG)
    draw = ImageDraw.Draw(img)
    rounded_box(draw, (40, 35, 1560, 865), CARD, outline="#D5DCE6", width=3, radius=34)
    rounded_box(draw, (70, 60, 1530, 170), NAVY, outline=None, width=0, radius=28)
    circle_label(draw, (140, 115), 42, theme, outline="#FFFFFF", text="", fnt=font(30, bold=True))
    # actual title
    draw_wrapped(draw, (230, 77), title, font(34, bold=True), "#FFFFFF", 1180, line_gap=0)
    if subtitle:
        draw_wrapped(draw, (230, 128), subtitle, font(22), "#E8EEF6", 1180, line_gap=0)
    return img, draw


def add_footer_note(draw, text):
    rounded_box(draw, (80, 790, 1520, 845), fill="#F3F6FB", outline="#D4DCE8", width=2, radius=18)
    draw.text((110, 804), text, font=font(20), fill=MUTED)


def make_step1_diagram(out: Path):
    img, draw = create_canvas(
        "Step 1. A city map becomes a living network",
        "The model turns Boston into tiny places joined by lines, so cooling flow can be studied like a road map.",
        theme=BLUE,
    )
    # left heatmap-like grid
    gx, gy = 110, 230
    cell = 54
    for r in range(6):
        for c in range(6):
            v = (r * 0.15 + c * 0.10) % 1.0
            col = (
                int(255 * (0.95 - 0.35 * v)),
                int(255 * (0.80 - 0.55 * v)),
                int(255 * (0.60 - 0.45 * v)),
            )
            draw.rectangle((gx + c * cell, gy + r * cell, gx + c * cell + cell - 4, gy + r * cell + cell - 4), fill=col)
    draw.text((120, 560), "Map cells", font=font(24, bold=True), fill=INK)
    draw.text((120, 595), "Each square is one place in the city.", font=font(20), fill=MUTED)
    arrow(draw, (460, 395), (650, 395), BLUE, width=10)
    draw.text((515, 345), "becomes", font=font(22, bold=True), fill=BLUE)
    # right network
    pts = [(780, 260), (920, 220), (1060, 285), (1220, 235), (1360, 310), (980, 470), (1150, 520), (1310, 455), (900, 575), (1095, 640)]
    edges = [(0,1),(1,2),(2,3),(3,4),(1,5),(5,6),(6,7),(5,8),(8,9),(2,6),(4,7)]
    for a, b in edges:
        draw.line([pts[a], pts[b]], fill="#90A4C2", width=7)
    for i, p in enumerate(pts):
        circle_label(draw, p, 24, fill=[TEAL, GREEN, GOLD, PURPLE, RED][i % 5], text=str(i + 1), fnt=font(18, bold=True))
    draw.text((820, 560), "Graph nodes and edges", font=font(24, bold=True), fill=INK)
    draw.text((820, 595), "Nearby places get linked so we can study who helps whom.", font=font(20), fill=MUTED)
    add_footer_note(draw, "Plain meaning: the city is not just a picture anymore. It becomes a network of connected places.")
    img.save(out)


def make_step2_diagram(out: Path):
    img, draw = create_canvas(
        "Step 2. Give every connection a strength",
        "A shaded, tree-lined link gets a stronger score than a hot, paved link.",
        theme=GREEN,
    )
    # three connection lanes
    y = 300
    starts = [180, 180, 180]
    ends = [620, 620, 620]
    colors = [GREEN, GOLD, RED]
    fills = [LIGHT_GREEN, LIGHT_GOLD, LIGHT_RED]
    labels = ["strong cooling link", "medium link", "weak barrier"]
    widths = [18, 10, 5]
    for i in range(3):
        rounded_box(draw, (120, y - 40, 680, y + 120), fill=fills[i], outline="#D6DEE8", width=2, radius=24)
        circle_label(draw, (220, y + 20), 34, fill=colors[i], text="A", fnt=font(26, bold=True))
        circle_label(draw, (540, y + 20), 34, fill=colors[i], text="B", fnt=font(26, bold=True))
        arrow(draw, (260, y + 20), (500, y + 20), colors[i], width=widths[i])
        draw.text((270, y - 10), labels[i], font=font(24, bold=True), fill=INK)
        y += 180
    rounded_box(draw, (760, 250, 1490, 690), fill="#F4F8FC", outline="#D6DEE8", width=2, radius=26)
    draw.text((800, 285), "Math snapshot", font=font(26, bold=True), fill=BLUE)
    draw_wrapped(
        draw,
        (800, 335),
        "w_ij = exp(-alpha * g_ij) * (1 + beta * ndvi_ij)",
        font(28, bold=False),
        INK,
        650,
        line_gap=8,
    )
    draw_wrapped(draw, (800, 430), "Big idea: more greenery and less barrier = a stronger cooling link.", font(24), MUTED, 650, line_gap=10)
    draw_wrapped(draw, (800, 540), "Boston lens: a link near parks, trees, or the river gets a better score than one across bare asphalt.", font(22), MUTED, 650, line_gap=8)
    add_footer_note(draw, "Plain meaning: not every neighborhood connection is equally helpful for cooling.")
    img.save(out)


def make_step3_diagram(out: Path):
    img, draw = create_canvas(
        "Step 3. Build the Laplacian, the city’s balance sheet",
        "This matrix keeps track of what is connected, and how much each place gives and receives.",
        theme=PURPLE,
    )
    # matrix block
    rounded_box(draw, (100, 240, 730, 720), fill="#F8F5FF", outline="#D7C7F2", width=2, radius=24)
    draw.text((140, 275), "Connections table", font=font(26, bold=True), fill=PURPLE)
    cell = 72
    for r in range(4):
        for c in range(4):
            x = 145 + c * cell
            y = 335 + r * cell
            fill = "#EADFFF" if r == c else "#FFFFFF"
            draw.rectangle((x, y, x + cell - 8, y + cell - 8), fill=fill, outline="#CDBBE8", width=2)
            if r == c:
                draw.text((x + 20, y + 18), "D", font=font(28, bold=True), fill=PURPLE)
            elif abs(r - c) == 1:
                draw.text((x + 18, y + 18), "W", font=font(28, bold=True), fill=INK)
    draw.text((145, 640), "D = total connection power", font=font(22), fill=MUTED)
    draw.text((145, 672), "W = who is linked to whom", font=font(22), fill=MUTED)
    # right-side formula and graph
    rounded_box(draw, (800, 240, 1490, 720), fill="#F6FBFF", outline="#CFE0F3", width=2, radius=24)
    draw.text((840, 275), "Math snapshot", font=font(26, bold=True), fill=BLUE)
    draw_wrapped(draw, (840, 340), "L = D - W", font(44, bold=True), INK, 560, line_gap=0)
    draw_wrapped(draw, (840, 430), "Normalized form: L_norm = I - D^{-1/2} W D^{-1/2}", font(24), MUTED, 620, line_gap=8)
    # balance graphic
    base_y = 620
    draw.line((930, base_y, 1370, base_y), fill="#6D7E95", width=8)
    draw.line((1150, 520, 1150, base_y), fill="#6D7E95", width=10)
    draw.polygon([(1150, 520), (1115, 575), (1185, 575)], fill="#6D7E95")
    rounded_box(draw, (980, 545, 1095, 600), fill=LIGHT_GREEN, outline=GREEN, width=3, radius=12)
    rounded_box(draw, (1205, 530, 1335, 605), fill=LIGHT_RED, outline=RED, width=3, radius=12)
    draw.text((998, 560), "more links", font=font(20, bold=True), fill=GREEN)
    draw.text((1218, 553), "less links", font=font(20, bold=True), fill=RED)
    draw.text((840, 690), "The matrix is just bookkeeping for the whole city at once.", font=font(22), fill=MUTED)
    add_footer_note(draw, "Plain meaning: the math asks, “how strongly does the city hold together?”")
    img.save(out)


def make_step6_diagram(out: Path):
    img, draw = create_canvas(
        "Step 6. Avoid impossible brute force",
        "The number of possible city splits grows too fast, so we use smart shortcuts.",
        theme=GOLD,
    )
    # explosion tree
    rounded_box(draw, (90, 235, 720, 735), fill="#FFF9EE", outline="#E7D6A5", width=2, radius=24)
    draw.text((130, 270), "Too many choices", font=font(26, bold=True), fill=GOLD)
    root = (220, 410)
    draw.ellipse((root[0]-22, root[1]-22, root[0]+22, root[1]+22), fill=GOLD)
    for i, dy in enumerate([-120, -40, 40, 120]):
        mid = (360, 410 + dy)
        draw.line([root, mid], fill="#B78A2E", width=7)
        draw.ellipse((mid[0]-18, mid[1]-18, mid[0]+18, mid[1]+18), fill="#E2B23E")
        for j, ddy in enumerate([-28, 0, 28]):
            leaf = (500, mid[1] + ddy)
            draw.line([mid, leaf], fill="#C59D3C", width=5)
            draw.ellipse((leaf[0]-13, leaf[1]-13, leaf[0]+13, leaf[1]+13), fill=GOLD)
    magnifier = (610, 500)
    draw.ellipse((magnifier[0]-50, magnifier[1]-50, magnifier[0]+50, magnifier[1]+50), outline=GOLD, width=10)
    draw.line((650, 540, 700, 595), fill=GOLD, width=12)
    draw.text((130, 625), "Exponential mess", font=font(22, bold=True), fill=RED)
    draw.text((130, 662), "Every extra node makes the search blow up.", font=font(20), fill=MUTED)
    # shortcut panel
    rounded_box(draw, (780, 235, 1490, 735), fill="#F6FBFF", outline="#D6DEE8", width=2, radius=24)
    draw.text((820, 270), "What we do instead", font=font(26, bold=True), fill=BLUE)
    steps = [
        ("Sweep", "rank nodes by one score", BLUE),
        ("Monte Carlo", "try random failures many times", TEAL),
        ("Greedy", "pick the best next move", GREEN),
    ]
    yy = 340
    for title, desc, color in steps:
        rounded_box(draw, (840, yy, 1415, yy + 100), fill="#FFFFFF", outline="#D6DEE8", width=2, radius=18)
        circle_label(draw, (890, yy + 50), 26, fill=color, text="1", fnt=font(20, bold=True))
        draw.text((935, yy + 26), title, font=font(24, bold=True), fill=INK)
        draw.text((935, yy + 58), desc, font=font(18), fill=MUTED)
        yy += 118
    draw.text((820, 690), "Plain meaning: we ask the smartest few questions instead of all possible ones.", font=font(21), fill=MUTED)
    add_footer_note(draw, "This is how the method stays fast enough to work on a real city like Boston.")
    img.save(out)


def make_step9_diagram(out: Path):
    img, draw = create_canvas(
        "Step 9. Smooth the city, then choose the best move",
        "A GMRF makes nearby places behave like neighbors, and the optimizer picks the most useful action first.",
        theme=TEAL,
    )
    # three-panel flow
    xs = [95, 555, 1015]
    titles = ["Prior", "Observations", "Posterior"]
    fills = ["#EEF8F6", "#F7FBFE", "#F1FBF6"]
    outlines = ["#B6E0D7", "#CFE0F3", "#BDE6D5"]
    for i in range(3):
        rounded_box(draw, (xs[i], 245, xs[i] + 360, 620), fill=fills[i], outline=outlines[i], width=2, radius=22)
        draw.text((xs[i] + 22, 268), titles[i], font=font(24, bold=True), fill=INK)
    # prior smooth field
    for r in range(5):
        for c in range(5):
            x = 135 + c * 58
            y = 330 + r * 58
            val = (r + c) / 8
            col = (
                int(235 - 40 * val),
                int(248 - 20 * val),
                int(245 - 15 * val),
            )
            draw.rectangle((x, y, x + 50, y + 50), fill=col, outline="#D6E8E2")
    draw.text((120, 580), "Nearby places stay similar.", font=font(18), fill=MUTED)
    # observations
    for r in range(5):
        for c in range(5):
            x = 595 + c * 58
            y = 330 + r * 58
            draw.rectangle((x, y, x + 50, y + 50), fill="#FFFFFF", outline="#D6E0EB")
    obs = [(625, 360), (740, 418), (684, 476), (800, 360)]
    for p in obs:
        circle_label(draw, p, 12, fill=RED, text="", fnt=font(12, bold=True))
    draw.text((580, 580), "A few measured points", font=font(18), fill=MUTED)
    # posterior
    for r in range(5):
        for c in range(5):
            x = 1055 + c * 58
            y = 330 + r * 58
            val = abs(2 - r) + abs(2 - c)
            col = (
                int(235 - 28 * (4 - min(val, 4)) / 4),
                int(248 - 38 * (4 - min(val, 4)) / 4),
                int(245 - 15 * (4 - min(val, 4)) / 4),
            )
            draw.rectangle((x, y, x + 50, y + 50), fill=col, outline="#D6E8E2")
    draw.text((1040, 580), "A smooth city-wide estimate", font=font(18), fill=MUTED)
    # arrows between panels
    arrow(draw, (455, 430), (525, 430), BLUE, width=8)
    arrow(draw, (915, 430), (985, 430), TEAL, width=8)
    # right-side optimization panel
    rounded_box(draw, (90, 660, 1490, 835), fill="#F4FBFA", outline="#CDEBE5", width=2, radius=20)
    draw.text((125, 684), "Math snapshot", font=font(24, bold=True), fill=TEAL)
    draw.text((125, 726), "Q = tau * L + eps * I    ->    Q_post = Q + sigma^-2 * I_obs", font=font(22), fill=INK)
    draw.text((125, 768), "Then the greedy search keeps the best next intervention, again and again.", font=font(20), fill=MUTED)
    add_footer_note(draw, "Plain meaning: the model does not just spot hot places. It turns them into the best action plan.")
    img.save(out)


def add_title_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("123B6D")
    return p


def set_style_font(style, name, size=None, bold=None, color=None):
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    if color is not None:
        style.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("5B6472")


def add_image(doc: Document, path: Path, width_inches: float = 6.4, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    if caption:
        add_caption(doc, caption)


def add_step_heading(doc: Document, title: str):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.color.rgb = RGBColor.from_string("2E74B5")
    return p


def add_body_paragraph(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r1 = p.add_run(f"{label} ")
    r1.bold = True
    r1.font.name = "Calibri"
    r1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string("1F2937")
    r2 = p.add_run(text)
    r2.font.name = "Calibri"
    r2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor.from_string("1F2937")
    return p


def add_formula_paragraph(doc: Document, formula: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(formula)
    r.font.name = "Consolas"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string("4B5563")
    return p


def build_doc():
    # Generate diagrams first.
    assets = {}
    make_step1_diagram(TMP_DIR / "step1_graph.png")
    make_step2_diagram(TMP_DIR / "step2_weights.png")
    make_step3_diagram(TMP_DIR / "step3_laplacian.png")
    make_step6_diagram(TMP_DIR / "step6_complexity.png")
    make_step9_diagram(TMP_DIR / "step9_gmrf.png")
    assets["step1"] = TMP_DIR / "step1_graph.png"
    assets["step2"] = TMP_DIR / "step2_weights.png"
    assets["step3"] = TMP_DIR / "step3_laplacian.png"
    assets["step6"] = TMP_DIR / "step6_complexity.png"
    assets["step9"] = TMP_DIR / "step9_gmrf.png"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    set_style_font(styles["Normal"], "Calibri", 11, False, "1F2937")
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    set_style_font(styles["Heading 1"], "Calibri", 16, False, "2E74B5")
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], "Calibri", 13, False, "2E74B5")
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], "Calibri", 12, False, "1F4D78")
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    add_title_paragraph(doc, "Urban Thermal Network Math Deep Dive")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("A picturized, layman-friendly guide to the math behind Boston cooling maps")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(12)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string("5B6472")
    doc.add_paragraph(
        "Based on docs/Urban_Thermal_Math_Deep_Dive.md, this version turns the nine main ideas into simple pictures, clear language, and Boston-oriented examples."
    )
    doc.add_paragraph(
        "The goal is not to water the math down. It is to make the math easy to see, easy to explain, and easy to use."
    )

    add_image(
        doc,
        ASSETS["boston_map"],
        width_inches=6.5,
        caption="Boston reference map from the repository. Brighter areas are hotter; cooler and greener areas act like cooling anchors.",
    )

    doc.add_heading("The 9-step story", level=1)
    doc.add_paragraph(
        "Read the model like a story: map -> graph -> weights -> Laplacian -> bottleneck -> robustness -> complexity -> cooling access -> priority plan."
    )

    # Step 1
    add_step_heading(doc, "Step 1. A city map becomes a living network")
    add_body_paragraph(
        doc,
        "Plain English:",
        "Boston is chopped into tiny places. Each place becomes a dot, and touching places get a line. That means heat, shade, parks, streets, and water can all be studied as one connected system.",
    )
    add_formula_paragraph(doc, "map cell -> node    |    touching cells -> edge    |    stronger link -> more cooling help")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "A park edge, the Charles River side, or a tree-lined street can behave like a cooling anchor because it gives the network an easier path for heat relief.",
    )
    add_image(doc, assets["step1"], width_inches=6.4, caption="Diagram: a grid of map cells turns into a graph of connected places.")

    # Step 2
    add_step_heading(doc, "Step 2. Give every connection a strength")
    add_body_paragraph(
        doc,
        "Plain English:",
        "Not every connection is equally useful. A shaded, greener path is stronger than a hot, paved path. The model scores each link so it can tell helpful connections from weak ones.",
    )
    add_formula_paragraph(doc, "w_ij = exp(-alpha * g_ij) * (1 + beta * ndvi_ij)")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "A route through trees, open space, or near water gets a better conductance score than a route crossing bare asphalt or big impermeable lots.",
    )
    add_image(doc, assets["step2"], width_inches=6.4, caption="Diagram: thick lines mean stronger cooling links; thin red lines mean weak barriers.")

    # Step 3
    add_step_heading(doc, "Step 3. Build the Laplacian, the city’s balance sheet")
    add_body_paragraph(
        doc,
        "Plain English:",
        "The Laplacian is bookkeeping. It compares how much connection each place has with who it is connected to. If the city were a team, this matrix would tell us whether the team is balanced or lopsided.",
    )
    add_formula_paragraph(doc, "L = D - W    and    L_norm = I - D^{-1/2} W D^{-1/2}")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "A place with many helpful neighbors has more balance. A place with few good neighbors stands out as structurally weaker in the thermal network.",
    )
    add_image(doc, assets["step3"], width_inches=6.4, caption="Diagram: the matrix records total connection power and subtracts the raw links.")

    # Step 4
    add_step_heading(doc, "Step 4. Find the narrowest bridge")
    add_body_paragraph(
        doc,
        "Plain English:",
        "The Cheeger cut looks for the thinnest part of the cooling network, like the smallest bridge between two big neighborhoods. If that bridge is weak, the city can split into two poorly connected halves.",
    )
    add_formula_paragraph(doc, "phi(S) = cut(S, V\\S) / min(vol(S), vol(V\\S))")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "If the bottleneck corridor lights up on the Boston map, that is where shade, trees, and other cooling improvements can have the biggest structural effect.",
    )
    add_image(doc, ASSETS["fig3"], width_inches=6.2, caption="Boston model output: the Cheeger thermal bottleneck priority map.")

    # Step 5
    add_step_heading(doc, "Step 5. Shake the city and see what breaks")
    add_body_paragraph(
        doc,
        "Plain English:",
        "The model randomly removes some links to see whether the network stays connected. This is a stress test. If the city falls apart too easily, the cooling system is fragile.",
    )
    add_formula_paragraph(doc, "each edge survives with probability p    ->    measure the giant connected piece")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "A network that only stays healthy when almost every link survives needs more backup paths, more canopy continuity, and fewer dead-end hot spots.",
    )
    add_image(doc, ASSETS["fig2"], width_inches=6.2, caption="Boston model output: bond percolation scan for connectivity robustness.")

    # Step 6
    add_step_heading(doc, "Step 6. Avoid impossible brute force")
    add_body_paragraph(
        doc,
        "Plain English:",
        "There are too many possible city splits to check one by one. So the model uses shortcuts: a spectral sweep, Monte Carlo simulation, and greedy choice instead of exhaustive search.",
    )
    add_formula_paragraph(doc, "too many subsets -> use sweep cuts, random sampling, and greedy selection")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "This is what makes the method practical on a full city instead of only on tiny toy examples. Boston-scale data would be too large for brute force.",
    )
    add_image(doc, assets["step6"], width_inches=6.4, caption="Diagram: combinatorial explosion on the left, fast approximation tools on the right.")

    # Step 7
    add_step_heading(doc, "Step 7. Measure how easy it is to reach cooling")
    add_body_paragraph(
        doc,
        "Plain English:",
        "The model asks a simple question: how far is this place from a cool, green, or watery refuge? It turns that distance into a cooling access score. High access is good; low access means heat is more trapped.",
    )
    add_formula_paragraph(doc, "r_ij ~ 1 / w_ij    and    access = 100 * (1 - normalized distance)")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "Parks, the Charles River edge, canopy-rich streets, and other cool places act like sinks. Neighborhoods far from them deserve more attention.",
    )
    add_image(doc, ASSETS["fig4"], width_inches=6.2, caption="Boston model output: cooling sink resistance proxy.")

    # Step 8
    add_step_heading(doc, "Step 8. Rank the best places to act first")
    add_body_paragraph(
        doc,
        "Plain English:",
        "Now the model blends heat and weak access into one urgency score. Hot places with poor cooling access rise to the top. That gives planners a clear first place to spend effort.",
    )
    add_formula_paragraph(doc, "priority = 100 * (0.65 * heat + 0.35 * (1 - access))")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "A hot downtown block with weak access to parks or trees will rank above a place that is hot but already well connected to cooling corridors.",
    )
    add_image(doc, ASSETS["fig1"], width_inches=6.2, caption="Boston model output: tradeoff between spectral gap and cooling reliability.")

    # Step 9
    add_step_heading(doc, "Step 9. Smooth the city, then choose the best move")
    add_body_paragraph(
        doc,
        "Plain English:",
        "A GMRF makes nearby places behave like neighbors instead of random strangers. Then the optimizer tries candidate interventions one at a time and keeps the one that improves the city the most.",
    )
    add_formula_paragraph(doc, "Q = tau * L + eps * I    ->    Q_post = Q + sigma^-2 * I_obs")
    add_body_paragraph(
        doc,
        "Boston lens:",
        "The final result is not just a hot-dot map. It is a connected plan for where shade, trees, cool pavement, or green links should go first.",
    )
    add_image(doc, assets["step9"], width_inches=6.4, caption="Diagram: prior knowledge, observations, and a smoothed posterior guide the final action plan.")

    doc.add_heading("Final takeaway", level=1)
    doc.add_paragraph(
        "The whole framework treats Boston like a connected cooling network. It finds where the network is weak, tests how fragile it is, measures access to cool places, and then ranks the best interventions. The math is serious, but the story is simple: connect hot places to cool places through the strongest possible corridors."
    )

    doc.add_paragraph(
        "Source used: docs/Urban_Thermal_Math_Deep_Dive.md and the repository’s Boston figures."
    )

    doc.save(str(OUT_DOCX))
    return OUT_DOCX


if __name__ == "__main__":
    out = build_doc()
    print(out)
