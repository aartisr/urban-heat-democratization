
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches
import zipfile
from pathlib import Path

def generate_pdf(pdf_path, figures=None, summary_text=None):
	c = canvas.Canvas(str(pdf_path), pagesize=letter)
	width, height = letter
	y = height - 50
	c.setFont("Helvetica-Bold", 16)
	c.drawString(50, y, "Spectral Urbanism Analysis Report")
	y -= 40
	c.setFont("Helvetica", 12)
	if summary_text:
		for line in summary_text.splitlines():
			c.drawString(50, y, line)
			y -= 18
	if figures:
		for fig in figures:
			if y < 200:
				c.showPage()
				y = height - 50
			c.drawString(50, y, f"Figure: {fig}")
			y -= 18
			try:
				c.drawImage(str(fig), 50, y-200, width=400, height=200, preserveAspectRatio=True, mask='auto')
				y -= 220
			except Exception:
				c.drawString(50, y, f"[Could not embed {fig}]")
				y -= 18
	c.save()

def generate_docx(docx_path, figures=None, summary_text=None):
	doc = Document()
	doc.add_heading("Spectral Urbanism Analysis Report", 0)
	if summary_text:
		doc.add_paragraph(summary_text)
	if figures:
		for fig in figures:
			doc.add_paragraph(f"Figure: {fig}")
			try:
				doc.add_picture(str(fig), width=Inches(5))
			except Exception:
				doc.add_paragraph(f"[Could not embed {fig}]")
	doc.save(str(docx_path))

def generate_zip(zip_path, files):
	with zipfile.ZipFile(str(zip_path), 'w') as zf:
		for f in files:
			f = Path(f)
			if f.exists():
				zf.write(str(f), arcname=f.name)
