from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=bold, color="1F2937")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="1F2937")


def build(out_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(4)
    r = title.add_run("Spectral Comparison")
    set_run_font(r, size=24, bold=True, color="123B6D")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Side-by-side comparison of `spectral-urbanism` and `spectral_urbanism_boston`.")
    set_run_font(r, size=11.5, italic=True, color="5B6472")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(
        "This document answers one question: is the Boston stack a full superset of the original app, or mainly a Boston-specific adaptation?"
    )
    set_run_font(r, size=11, color="1F2937")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    cell_text(hdr[0], "Feature", bold=True)
    cell_text(hdr[1], "spectral-urbanism", bold=True)
    cell_text(hdr[2], "spectral_urbanism_boston", bold=True)

    rows = [
        ("Core idea", "Raster heat/NDVI maps turned into a graph network.", "City-specific thermal network built for Boston.", "Same core idea."),
        ("Graph model", "Weighted grid graph from rasters.", "Graph over city grid cells / GeoDataFrame rows, plus touch adjacency and optional wind edges.", "Boston is more city-specific."),
        ("Spectral gap / λ₂", "Yes.", "Yes, documented.", "Likely yes."),
        ("Cheeger bottleneck", "Yes.", "Yes, documented.", "Likely yes."),
        ("Resistance / cooling access", "Yes.", "Yes, documented.", "Likely yes."),
        ("Percolation / robustness", "Yes.", "Yes.", "Likely yes."),
        ("Reliability", "Sink-reachability reliability.", "All-terminal reliability in the Boston docs.", "Similar, but not identical."),
        ("GMRF inference", "Not clearly visible in this repo’s app code.", "Explicitly documented.", "Boston adds this."),
        ("Greedy optimization", "Yes, intervention scenario in the pipeline.", "Yes, documented.", "Boston adds a clearer formulation."),
        ("UI pages", "Full Streamlit app: Theory, Comparison, Results, Export, Data.", "No Boston app UI is present in this workspace.", "Original app is more complete."),
        ("API/backend", "FastAPI backend present.", "Not verified here.", "Original app is more complete."),
        ("Exports", "PDF, DOCX, ZIP, KML, GeoJSON, GeoTIFF.", "Not verified here.", "Original app has broader verified exports."),
    ]

    for feature, left, right, verdict in rows:
        row = table.add_row().cells
        cell_text(row[0], feature, bold=True)
        cell_text(row[1], left)
        cell_text(row[2], right)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Bottom line")
    set_run_font(r, size=14, bold=True, color="2E74B5")

    add_bullet(doc, "Boston appears to include the main scientific machinery: graph construction, spectral analysis, Cheeger cuts, robustness, resistance/access, GMRF, and optimization.")
    add_bullet(doc, "The original `spectral-urbanism` repo is the more complete end-to-end application in this workspace because its UI, API, export pipeline, and report generation are directly visible in code.")
    add_bullet(doc, "So the safest answer is: Boston is a strong Boston-specific adaptation, but not a proven 1:1 superset of every feature from the original app.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Source basis: ")
    set_run_font(r, size=10.5, bold=True, color="1F2937")
    r = p.add_run("README.md, core/pipeline.py, pages/Comparison.py, pages/Results.py, and docs/Urban_Thermal_Math_Deep_Dive.md.")
    set_run_font(r, size=10.5, color="1F2937")

    doc.save(str(out_path))


if __name__ == "__main__":
    out = Path("/Users/rraviku2/aarti/spectral_urbanism_boston/docs/spectral_comparison.docx")
    build(out)
    print(out)
